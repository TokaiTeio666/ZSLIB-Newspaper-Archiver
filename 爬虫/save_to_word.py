import time
import re
import io
from docx import Document
from docx.shared import Inches
from selenium.webdriver.common.by import By
from capture import capture
import os

def save_to_separate_word(name, info, detail_url, driver,row_index, output_dir):
    """
    针对单条记录生成高清裁剪图并保存为独立 Word 文档
    """
    try:
        # 1. 记录当前列表页句柄
        list_window = driver.current_window_handle

        # 2. 调用 capture 函数：内部会执行 driver.get(detail_url) 并返回 PIL Image 对象
        # 这里的裁剪逻辑基于您提供的 capture.py
        cropped_image = capture(detail_url, driver)

        # 3. 创建新文档并写入文本信息
        doc = Document()
        date_parts = info.split()
        if len(date_parts) >= 3:
            year, month, day = date_parts[0], date_parts[1], date_parts[2]
            edition = date_parts[3] if len(date_parts) > 3 else ""

            # 按照要求的格式分行写入日期
            doc.add_paragraph(f"{year}年")
            doc.add_paragraph(f"{month}月")
            doc.add_paragraph(f"{day}日")

            # 4. 将 PIL 图像对象转换为字节流插入文档，无需保存临时文件
            img_byte_arr = io.BytesIO()
            cropped_image.save(img_byte_arr, format='PNG')
            doc.add_picture(img_byte_arr, width=Inches(5))

            # 5. 写入来源标注
            source_text = f"——《{name}》{year}年{month}月{day}日，{edition}"
            doc.add_paragraph(source_text)

            # 6. 保存独立文档 (清洗文件名)
            clean_name = re.sub(r'[\\/:*?"<>|]', '_', name)
            file_name = f"{year}年{month}月{day}日_{clean_name}.docx"
            final_save_path = os.path.join(output_dir, file_name)
            doc.save(final_save_path)
            print(f"成功保存高清文档: {file_name}")

    except Exception as e:
        print(f"处理 {name} 时出错: {e}")
    finally:
        # 核心：处理完详情页后，必须确保回到搜索列表页，否则后续循环找不到元素
        # 由于 capture 内部用了 driver.get，我们可能需要重新加载列表页或后退
        driver.back()
        time.sleep(2)
        # 重新切入 iframe，因为页面刷新或回退后 iframe 上下文会丢失
        try:
            target_frame = driver.find_elements(By.ID, "BriefList")[1]
            driver.switch_to.frame(target_frame)
        except:
            pass