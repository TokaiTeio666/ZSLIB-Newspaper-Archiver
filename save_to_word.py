import time
import re
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from selenium.webdriver.common.by import By
from capture import capture, ocr_text
import os


def save_to_separate_word(name, info, detail_url, driver, row_index, output_dir, article_title=""):
    """
    针对单条记录生成高清裁剪图并保存为独立 Word 文档
    文档格式：
      xx年
      xx月
      xx日
      文字内容（篇名、作者等）
      图片
      ——《xx日报》xx年xx月xx日，第xx版

    注意：detail_url 参数已废弃，直接使用当前driver（已在详情页）
    """
    try:
        # 1. 调用 capture 函数（直接使用当前driver，不需要url）
        # 新版 capture 返回 (PIL Image, text_info字典)
        result = capture("", driver, output_dir)

        # 兼容处理返回值
        if isinstance(result, tuple):
            cropped_image, text_info = result
        else:
            cropped_image = result
            text_info = {"title": "", "author": "", "content": "", "full_text": ""}

        # 如果capture返回None，跳过
        if cropped_image is None:
            print(f"警告：{name} 未能获取到图片，跳过保存。")
            return

        # 3. 解析日期和版次信息
        year, month, day, edition = parse_date_info(info, name, article_title)

        # 4. 创建新文档
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)

        # 5. 写入日期（分行）
        if year:
            doc.add_paragraph(f"{year}年")
        if month:
            doc.add_paragraph(f"{month}月")
        if day:
            doc.add_paragraph(f"{day}日")

        # 6. 写入文字内容（篇名、作者、正文等）
        # 可选 OCR 增强：未实现/未安装依赖时静默返回空串
        ocr_result = ""
        try:
            ocr_result = ocr_text(cropped_image)
        except Exception:
            ocr_result = ""
        article_text = build_article_text(text_info, article_title, ocr_result)
        if article_text:
            # 文字内容段落
            text_para = doc.add_paragraph(article_text)
            text_para.paragraph_format.space_after = Pt(6)

        # 7. 插入图片
        if cropped_image:
            img_byte_arr = io.BytesIO()
            cropped_image.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)
            try:
                doc.add_picture(img_byte_arr, width=Inches(5.5))
                # 图片居中
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"插入图片失败: {e}")
                doc.add_paragraph("[图片插入失败]")

        # 8. 写入来源标注
        source_parts = []
        if name:
            source_parts.append(f"《{name}》")
        if year and month and day:
            source_parts.append(f"{year}年{month}月{day}日")
        if edition:
            source_parts.append(f"，{edition}")

        source_text = "——" + "".join(source_parts)
        source_para = doc.add_paragraph(source_text)
        source_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 9. 保存文档（清洗文件名）
        clean_name = re.sub(r'[\\/:*?"<>|]', '_', name) if name else "未命名"
        date_prefix = ""
        if year and month and day:
            date_prefix = f"{year}年{month}月{day}日_"
        file_name = f"{date_prefix}{clean_name}.docx"
        final_save_path = os.path.join(output_dir, file_name)

        # 处理重名
        counter = 1
        base_path = final_save_path
        while os.path.exists(final_save_path):
            final_save_path = os.path.join(
                output_dir,
                f"{date_prefix}{clean_name}_{counter}.docx"
            )
            counter += 1

        doc.save(final_save_path)
        print(f"成功保存文档: {os.path.basename(final_save_path)}")

    except Exception as e:
        print(f"处理 {name} 时出错: {e}")
        import traceback
        traceback.print_exc()
    finally:
        # 确保回到搜索列表页
        try:
            driver.back()
            time.sleep(2)
        except Exception:
            pass

        # 如果回退不成功，尝试切换回原窗口
        try:
            if driver.current_window_handle != list_window:
                driver.switch_to.window(list_window)
                time.sleep(1)
        except Exception:
            pass

        # 重新切入结果列表iframe（如果存在）
        try:
            target_frames = driver.find_elements(By.ID, "BriefList")
            if len(target_frames) > 1:
                driver.switch_to.frame(target_frames[1])
        except Exception:
            pass


def parse_date_info(info, name="", article_title=""):
    """
    从信息字符串中解析年、月、日、版次
    支持多种格式：
      - "1935年11月19日 第3版"
      - "1935-11-19 3版"
      - "1935/11/19 第3版"
      - "1935 11 19 3版"
    """
    year = month = day = edition = ""

    # 合并所有文本进行搜索
    full_text = f"{info} {name} {article_title}"

    # 匹配日期模式
    # 模式1: 1935年11月19日
    match = re.search(r"(\d{4})年(\d{1,2})月(\d{1,2})日", full_text)
    if match:
        year, month, day = match.group(1), match.group(2), match.group(3)
    else:
        # 模式2: 1935-11-19 或 1935/11/19
        match = re.search(r"(\d{4})[-/](\d{1,2})[-/](\d{1,2})", full_text)
        if match:
            year, month, day = match.group(1), match.group(2), match.group(3)
        else:
            # 模式3: 空格分隔 1935 11 19
            match = re.search(r"(\d{4})\s+(\d{1,2})\s+(\d{1,2})", full_text)
            if match:
                year, month, day = match.group(1), match.group(2), match.group(3)

    # 匹配版次
    edition_match = re.search(r"第?(\d+)\s*版", full_text)
    if edition_match:
        edition_num = edition_match.group(1)
        edition = f"第{edition_num}版"
    else:
        # 尝试匹配 "版次:3" 或 "版次：3"
        edition_match2 = re.search(r"版次[：:\s]*(\d+)", full_text)
        if edition_match2:
            edition = f"第{edition_match2.group(1)}版"

    return year, month, day, edition


def build_article_text(text_info, article_title="", ocr_result=""):
    """
    根据文字信息构建要写入文档的文字内容
    优先使用 article_title（搜索结果中的篇名），然后是 text_info 中的内容，
    最后补充 OCR 识别正文（可选增强，默认空串）
    """
    parts = []

    # 篇名/标题
    title = article_title or text_info.get("title", "")
    if title:
        parts.append(title)

    # 作者
    author = text_info.get("author", "")
    if author:
        parts.append(f"作者：{author}")

    # 正文内容
    content = text_info.get("content", "")
    if content and content != title:
        parts.append(content)

    # 可选 OCR 识别正文
    if ocr_result and ocr_result.strip():
        parts.append(ocr_result.strip())

    # 如果以上都没有，使用full_text
    if not parts:
        full_text = text_info.get("full_text", "").strip()
        if full_text:
            parts.append(full_text)

    return "\n".join(parts)
